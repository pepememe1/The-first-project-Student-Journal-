"""
exports.py — экспорт ведомости и журнала в xlsx/docx на ДЕСКТОПЕ.

Единый стиль десктопа и веба (требование заказчика): весь текст Times New Roman 14,
БЕЗ цветов (чёрный на белом), тонкие рамки, адаптивная ширина. Функции чистые — на
вход dict-строки (без ORM/Qt), поэтому легко тестируются и совпадают 1:1 с серверными
(server/app/xlsx_export.py, docx_export.py); здесь — ведомость (xlsx+docx) и журнал в docx.
⚠️ Раньше тут значилось «журнал в xlsx строит сам GradeBook (core.export_to_excel)» —
и класс, и его экспорт удалены 31.08.2026 как мёртвые (ноль вызывающих в продукте).

python-docx импортируется ЛЕНИВО внутри docx-функций: без пакета xlsx-путь работает,
а Word-экспорт понятно сообщит, что нужен пакет.
"""
import io
import re
from datetime import datetime

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter

FNT = "Times New Roman"
SZ = 14
_THIN = Side(style="thin", color="000000")
BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)


def _sem_txt(sem) -> str:
    return "осенний" if sem == 1 else ("весенний" if sem == 2 else str(sem))


def _fio(s: dict) -> str:
    """ФИО из строки студента; если отчество уже «зашито» в имя — не дублируем."""
    fio = " ".join(x for x in (s.get("surname", ""), s.get("name", ""),
                               s.get("patronymic", "")) if x).strip()
    if s.get("patronymic") and s.get("name", "").endswith(s.get("patronymic", "")):
        fio = f"{s.get('surname', '')} {s.get('name', '')}".strip()
    return fio


# ── XLSX ────────────────────────────────────────────────────────────────────────
def _autofit(ws, ncols: int, rows: list, min_w=8, max_w=48):
    for c in range(1, ncols + 1):
        best = 0
        for r in rows:
            v = ws.cell(row=r, column=c).value
            if v is None:
                continue
            best = max(best, max((len(s) for s in str(v).split("\n")), default=0))
        ws.column_dimensions[get_column_letter(c)].width = max(min_w, min(max_w, best * 1.45 + 2))


def _title(ws, row, last_col, text, size, bold=False, italic=False):
    ws.merge_cells(f"A{row}:{last_col}{row}")
    c = ws[f"A{row}"]
    c.value = text
    c.font = Font(name=FNT, size=size, bold=bold, italic=italic)
    c.alignment = Alignment(horizontal="center", vertical="center")


def build_vedomost_xlsx(group: str, subject: str, term: dict, form: str, rows,
                        teacher: str = "") -> bytes:
    """Экзаменационно-зачётная (аттестационная) ведомость. rows — [{surname, name,
    patronymic, grade}]. term — {year, semester}."""
    year = (term or {}).get("year", "")
    sem = (term or {}).get("semester", "")

    wb = Workbook()
    ws = wb.active
    ws.title = re.sub(r'[\[\]:*?/\\]', '_', f"Ведомость {group}")[:31]
    headers = ["№", "Фамилия Имя Отчество", "Итоговая оценка", "Подпись"]
    ncols = len(headers)
    last_col = get_column_letter(ncols)

    _title(ws, 1, last_col, "Технологический колледж ВСГУТУ", 12)
    _title(ws, 2, last_col, "Экзаменационная ведомость" if (form or "").lower().startswith("экз")
           else "Зачётно-экзаменационная ведомость", 16, bold=True)
    _title(ws, 3, last_col, f"Группа {group}  ·  {subject}", SZ, bold=True)
    _title(ws, 4, last_col, f"{year}, {_sem_txt(sem)} семестр"
           + (f"  ·  форма контроля: {form}" if form else ""), 12, italic=True)

    HDR = 6
    for _ in range(5):
        ws.append([])
    ws.append(headers)
    for cell in ws[HDR]:
        cell.font = Font(name=FNT, size=SZ, bold=True)
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = BORDER

    for i, s in enumerate(rows):
        grade = (s.get("grade") or "").strip()
        ws.append([i + 1, _fio(s), grade, ""])
        r = HDR + 1 + i
        for c in range(1, ncols + 1):
            cell = ws.cell(row=r, column=c)
            cell.border = BORDER
            cell.font = Font(name=FNT, size=SZ)
            cell.alignment = Alignment(horizontal="left" if c == 2 else "center", vertical="center")

    foot = HDR + len(rows) + 2
    ws.cell(row=foot, column=1, value=f"Дата: {datetime.now().strftime('%d.%m.%Y')}").font = Font(name=FNT, size=SZ)
    ws.merge_cells(start_row=foot, start_column=1, end_row=foot, end_column=2)
    sign = ws.cell(row=foot, column=3, value=f"Преподаватель: {teacher or '____________'}")
    sign.font = Font(name=FNT, size=SZ)
    ws.merge_cells(start_row=foot, start_column=3, end_row=foot, end_column=ncols)

    _autofit(ws, ncols, [HDR] + list(range(HDR + 1, HDR + len(rows) + 1)), min_w=6, max_w=55)
    ws.column_dimensions["A"].width = 5
    ws.row_dimensions[HDR].height = 30
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── DOCX (python-docx — лениво) ───────────────────────────────────────────────────
def build_journal_docx(group: str, subject: str, lessons, rows) -> bytes:
    """Журнал в Word (альбомная). lessons — объекты Lesson (id/type/number/topic/date/
    retake_date); rows — [{surname, name, records, average}]."""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT

    _ALIGN = {"center": WD_ALIGN_PARAGRAPH.CENTER, "left": WD_ALIGN_PARAGRAPH.LEFT,
              "right": WD_ALIGN_PARAGRAPH.RIGHT}

    def _cell(cell, text, bold=False, size=SZ, align="center"):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = _ALIGN[align]
        run = p.add_run("" if text is None else str(text))
        run.font.name = FNT
        run.font.size = Pt(size)
        run.bold = bold

    def _ttl(doc, text, size, bold=False, italic=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.name = FNT
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Cm(1.0))

    _ttl(doc, "Журнал успеваемости", 16, bold=True)
    _ttl(doc, f"Группа {group}  ·  {subject}", SZ, bold=True)
    _ttl(doc, f"Выгружено {datetime.now().strftime('%d.%m.%Y %H:%M')}  ·  "
              f"Технологический колледж ВСГУТУ", 12, italic=True)

    headers = ["Фамилия", "Имя"]
    keys = []
    for l in lessons:
        if l.type == "Экзамен":
            headers.append(f"Экзамен №{l.number} ({l.date}) {l.topic or ''}".strip())
            keys.append(l.id)
            if getattr(l, "retake_date", ""):
                headers.append(f"Пересдача ({l.retake_date})")
                keys.append(l.id + "_retake")
        else:
            headers.append(f"{l.type} №{l.number} ({l.date})")
            keys.append(l.id)
    headers.append("Средний балл")
    ncols = len(headers)

    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _cell(hdr[i], h, bold=True)

    for s in rows:
        recs = s.get("records") or {}
        cells = table.add_row().cells
        _cell(cells[0], s.get("surname", ""), align="left")
        _cell(cells[1], s.get("name", ""), align="left")
        for j, k in enumerate(keys):
            val = recs.get(k, "")
            if not val and k.endswith("_retake"):
                base = (recs.get(k[: -len("_retake")], "") or "").strip()
                failed = base.startswith(("2", "Н")) or "Не зачтено" in base
                val = "" if failed else "—"
            _cell(cells[2 + j], val or "·")
        avg = round(float(s.get("average") or 0), 2)
        _cell(cells[ncols - 1], avg if avg > 0 else "—", bold=True)

    vals = [round(float(s.get("average") or 0), 2) for s in rows]
    vals = [v for v in vals if v > 0]
    p = doc.add_paragraph()
    r = p.add_run(f"Средний по группе: {round(sum(vals) / len(vals), 2) if vals else '—'}")
    r.font.name = FNT
    r.bold = True
    r.font.size = Pt(SZ)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_vedomost_docx(group: str, subject: str, term: dict, form: str, rows,
                        teacher: str = "") -> bytes:
    """Ведомость в Word (книжная). rows — [{surname, name, patronymic, grade}]."""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    _ALIGN = {"center": WD_ALIGN_PARAGRAPH.CENTER, "left": WD_ALIGN_PARAGRAPH.LEFT,
              "right": WD_ALIGN_PARAGRAPH.RIGHT}

    def _cell(cell, text, bold=False, size=SZ, align="center"):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = _ALIGN[align]
        run = p.add_run("" if text is None else str(text))
        run.font.name = FNT
        run.font.size = Pt(size)
        run.bold = bold

    def _ttl(doc, text, size, bold=False, italic=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.name = FNT
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic

    year = (term or {}).get("year", "")
    sem = (term or {}).get("semester", "")

    doc = Document()
    for m in ("left_margin", "right_margin"):
        setattr(doc.sections[0], m, Cm(1.8))
    _ttl(doc, "Технологический колледж ВСГУТУ", 12)
    _ttl(doc, "Экзаменационная ведомость" if (form or "").lower().startswith("экз")
         else "Зачётно-экзаменационная ведомость", 16, bold=True)
    _ttl(doc, f"Группа {group}  ·  {subject}", SZ, bold=True)
    _ttl(doc, f"{year}, {_sem_txt(sem)} семестр" + (f"  ·  форма контроля: {form}" if form else ""),
         12, italic=True)
    doc.add_paragraph()

    headers = ["№", "Фамилия Имя Отчество", "Итоговая оценка", "Подпись"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _cell(hdr[i], h, bold=True)

    for i, s in enumerate(rows):
        cells = table.add_row().cells
        _cell(cells[0], i + 1)
        _cell(cells[1], _fio(s), align="left")
        _cell(cells[2], (s.get("grade") or "").strip())
        _cell(cells[3], "")

    doc.add_paragraph()
    foot = doc.add_paragraph()
    run = foot.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y')}"
                       f"          Преподаватель: {teacher or '____________'}")
    run.font.name = FNT
    run.font.size = Pt(SZ)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
