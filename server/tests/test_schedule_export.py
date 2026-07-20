"""
test_schedule_export.py — выгрузка расписания группы в xlsx и docx.

Импорта расписания у нас нет (нет и файла-образца от ВСГУТУ), поэтому делаем обратное:
отдаём разобранное с портала расписание файлом.

Главное, что здесь защищается: выгрузка строится из ТОГО ЖЕ слитого расписания
(портал + правки админа), которое видит студент на сайте. Разойдись они — и файлу
перестанут доверять, а бумажное расписание в колледже висит на стене.

Портал в тестах недоступен, поэтому расписание целиком складывается из админ-правок —
это же боевой сценарий «колледж без портала».
"""
import io

from conftest import make_admin


def _set(client, admin, **kw):
    r = client.post("/web/admin/schedule/override", json=kw, headers=admin)
    assert r.status_code == 200, r.text


def _seed_week(client, admin, group="ИС-21"):
    _set(client, admin, group=group, week=1, day="Пнд", pair_no=2, subject="Математика",
         teacher="Иванов И.И.", room="305", time="10:45-12:20", kind="Лекция")
    _set(client, admin, group=group, week=1, day="Срд", pair_no=1, subject="Физика",
         teacher="Петров П.П.", room="101", time="09:00-10:35")


def test_xlsx_contains_subject_teacher_and_room(client):
    from openpyxl import load_workbook
    admin = make_admin(client)
    _seed_week(client, admin)

    r = client.get("/web/schedule/export", params={"group": "ИС-21", "fmt": "xlsx"},
                   headers=admin)
    assert r.status_code == 200, r.text
    assert r.headers["content-type"].startswith("application/vnd.openxmlformats")

    ws = load_workbook(io.BytesIO(r.content)).active
    text = "\n".join(str(c.value) for row in ws.iter_rows() for c in row if c.value)
    for need in ("ИС-21", "Математика", "Иванов И.И.", "305", "Физика", "10:45-12:20"):
        assert need in text, f"в выгрузке нет: {need}"
    #Дни недели — это колонки сетки, без них таблица не читается.
    for day in ("Пнд", "Срд"):
        assert day in text


def test_docx_contains_the_same_data(client):
    from docx import Document
    admin = make_admin(client)
    _seed_week(client, admin)

    r = client.get("/web/schedule/export", params={"group": "ИС-21", "fmt": "docx"},
                   headers=admin)
    assert r.status_code == 200, r.text
    doc = Document(io.BytesIO(r.content))
    text = "\n".join(c.text for t in doc.tables for row in t.rows for c in row.cells)
    text += "\n" + "\n".join(p.text for p in doc.paragraphs)
    for need in ("ИС-21", "Математика", "Иванов И.И.", "305", "Физика"):
        assert need in text, f"в выгрузке нет: {need}"


def test_export_follows_admin_edits(client):
    """Правку админа выгрузка обязана показывать: файл строится из слитого расписания,
    а не из голого портала."""
    from openpyxl import load_workbook
    admin = make_admin(client)
    _seed_week(client, admin)
    #Заменяем пару в той же ячейке — в файле должен оказаться НОВЫЙ предмет.
    _set(client, admin, group="ИС-21", week=1, day="Пнд", pair_no=2, subject="Химия",
         teacher="Сидоров С.С.", room="404", time="10:45-12:20")

    r = client.get("/web/schedule/export", params={"group": "ИС-21"}, headers=admin)
    ws = load_workbook(io.BytesIO(r.content)).active
    text = "\n".join(str(c.value) for row in ws.iter_rows() for c in row if c.value)
    assert "Химия" in text and "404" in text
    assert "Математика" not in text, "заменённая пара не должна остаться в выгрузке"


def test_hidden_pair_is_not_exported(client):
    """action=remove скрывает пару на сайте — значит и в файле её быть не должно."""
    from openpyxl import load_workbook
    admin = make_admin(client)
    _seed_week(client, admin)
    _set(client, admin, group="ИС-21", week=1, day="Срд", pair_no=1, action="remove")

    r = client.get("/web/schedule/export", params={"group": "ИС-21"}, headers=admin)
    ws = load_workbook(io.BytesIO(r.content)).active
    text = "\n".join(str(c.value) for row in ws.iter_rows() for c in row if c.value)
    assert "Физика" not in text
    assert "Математика" in text, "остальные пары должны остаться"


def test_missing_schedule_gives_404(client):
    admin = make_admin(client)
    r = client.get("/web/schedule/export", params={"group": "НЕТ-ТАКОЙ"}, headers=admin)
    assert r.status_code == 404


def test_export_requires_auth(client):
    assert client.get("/web/schedule/export", params={"group": "ИС-21"}).status_code == 401
