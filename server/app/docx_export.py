"""
docx_export.py — сборка журнала и ведомости в Word (.docx).

Единый стиль с xlsx и десктопом: Times New Roman 14, БЕЗ цветов (чёрный текст),
таблицы с рамками (Table Grid). Журнал — АЛЬБОМНАЯ ориентация (много колонок),
ведомость — книжная. Word подгоняет ширину сам — текст не обрезается.
"""
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

FNT = "Times New Roman"
SZ = 14
_ALIGN = {"center": WD_ALIGN_PARAGRAPH.CENTER, "left": WD_ALIGN_PARAGRAPH.LEFT,
          "right": WD_ALIGN_PARAGRAPH.RIGHT}


def _cell(cell, text, bold=False, size=SZ, align="center"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = _ALIGN[align]
    run = p.add_run("" if text is None else str(text))
    run.font.name = FNT
    run.font.size = Pt(size)
    run.bold = bold


def _title(doc, text, size, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = FNT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _landscape(doc):
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Cm(1.0))


def build_journal_docx(group: str, subject: str, lessons, rows) -> bytes:
    """rows — [{surname, name, records: {lesson_id→оценка}, average}]. Как в xlsx-версии."""
    doc = Document()
    _landscape(doc)
    _title(doc, "Журнал успеваемости", 16, bold=True)
    _title(doc, f"Группа {group}  ·  {subject}", SZ, bold=True)
    _title(doc, f"Выгружено {datetime.now().strftime('%d.%m.%Y %H:%M')}  ·  "
                f"Технологический колледж ВСГУТУ", 12, italic=True)

    headers = ["Фамилия", "Имя"]
    keys = []
    for l in lessons:
        if l.type == "Экзамен":
            headers.append(f"Экзамен №{l.number} ({l.date}) {l.topic or ''}".strip())
            keys.append(l.id)
            if l.retake_date:
                headers.append(f"Пересдача ({l.retake_date})")
                keys.append(l.id + "_retake")
        else:
            headers.append(f"{l.type} №{l.number} ({l.date})")
            keys.append(l.id)
    headers.append("Средний балл")
    ncols = len(headers)

    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _cell(hdr[i], h, bold=True)

    for s in rows:
        recs = s.get("records") or {}
        cells = table.add_row().cells
        _cell(cells[0], s.get("surname", ""), align="left")
        _cell(cells[1], s.get("name", ""), align="left")
        for j, k in enumerate(keys):
            val = recs.get(k, "")
            if not val and k.endswith("_retake"):
                base = (recs.get(k[: -len("_retake")], "") or "").strip()
                failed = base.startswith(("2", "Н")) or "Не зачтено" in base
                val = "" if failed else "—"
            _cell(cells[2 + j], val or "·")
        avg = round(float(s.get("average") or 0), 2)
        _cell(cells[ncols - 1], avg if avg > 0 else "—", bold=True)

    vals = [round(float(s.get("average") or 0), 2) for s in rows]
    vals = [v for v in vals if v > 0]
    p = doc.add_paragraph()
    r = p.add_run(f"Средний по группе: {round(sum(vals) / len(vals), 2) if vals else '—'}")
    r.font.name = FNT
    r.bold = True
    r.font.size = Pt(SZ)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_vedomost_docx(group: str, subject: str, term: dict, form: str, rows,
                        teacher: str = "") -> bytes:
    """rows — [{surname, name, patronymic, grade}]. term — {year, semester}."""
    year = (term or {}).get("year", "")
    sem = (term or {}).get("semester", "")
    sem_txt = "осенний" if sem == 1 else ("весенний" if sem == 2 else str(sem))

    doc = Document()
    for m in ("left_margin", "right_margin"):
        setattr(doc.sections[0], m, Cm(1.8))
    _title(doc, "Технологический колледж ВСГУТУ", 12)
    _title(doc, "Экзаменационная ведомость" if (form or "").lower().startswith("экз")
           else "Зачётно-экзаменационная ведомость", 16, bold=True)
    _title(doc, f"Группа {group}  ·  {subject}", SZ, bold=True)
    _title(doc, f"{year}, {sem_txt} семестр" + (f"  ·  форма контроля: {form}" if form else ""),
           12, italic=True)
    doc.add_paragraph()

    headers = ["№", "Фамилия Имя Отчество", "Итоговая оценка", "Подпись"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _cell(hdr[i], h, bold=True)

    for i, s in enumerate(rows):
        fio = " ".join(x for x in (s.get("surname", ""), s.get("name", ""),
                                   s.get("patronymic", "")) if x).strip()
        if s.get("patronymic") and s.get("name", "").endswith(s.get("patronymic", "")):
            fio = f"{s.get('surname','')} {s.get('name','')}".strip()
        cells = table.add_row().cells
        _cell(cells[0], i + 1)
        _cell(cells[1], fio, align="left")
        _cell(cells[2], (s.get("grade") or "").strip())
        _cell(cells[3], "")

    doc.add_paragraph()
    foot = doc.add_paragraph()
    run = foot.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y')}"
                       f"          Преподаватель: {teacher or '____________'}")
    run.font.name = FNT
    run.font.size = Pt(SZ)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
