"""
curator_report.py — сводный отчёт куратора по успеваемости групп (Excel / Word).

Что внутри отчёта по КАЖДОЙ группе:
  • список студентов (ФИО по алфавиту);
  • средний балл по каждому предмету + общий средний;
  • посещаемость: Н (неявки), Б (болезнь), О (уважит.) и всего пропусков;
  • число долгов (несданные практики/экзамены).
Это покрывает три вещи, которые просил заказчик: аналитику посещения, оценки и списки
групп — в одном документе, с выбором формата (xlsx/docx) и охвата (одна/несколько/все
курируемые группы).

Данные берём ТОЛЬКО через webdata (те же выборки, что видит куратор на экране), поэтому
отчёт не расходится с интерфейсом. Средний — единый grading.practice_average.
"""
import io

from . import webdata as W


def parse_ddmmyyyy(s: str):
    """"ДД.ММ.ГГГГ" → кортеж (год,месяц,день) для сравнения, или None — не распознали.
    Формат дат занятий именно такой (см. teacher_create_lesson)."""
    try:
        d, m, y = (s or "").strip().split(".")
        return (int(y), int(m), int(d))
    except Exception:
        return None


def collect_group(db, group: str, year: str, semester: int, cfg: dict,
                  cutoff_date: str = "") -> dict:
    """Сводка по одной группе: студенты × (средние по предметам + посещаемость + долги).

    `cutoff_date` (опц., "ДД.ММ.ГГГГ") — снимок «на дату включительно»: занятия ПОСЛЕ
    неё не учитываются. Нужен отчёту куратора в мессенджере (§12 плана) — кнопка
    «Отчёт №X» должна показывать состояние на момент СОЗДАНИЯ, а не текущее (см.
    routers/messenger.py::_handle_report_command). Без cutoff_date — как раньше,
    весь термин целиком (экспорт xlsx/docx куратора этим путём и продолжает ходить)."""
    students = W.students_in_group(db, group)
    lessons = W.group_lessons(db, group, year=year, semester=semester)
    #⚠️ Живой баг («отчёты показывают старые предметы, а не актуальные»): список
    #предметов раньше собирался ПРЯМО ИЗ ЗАНЯТИЙ, а удаление предмета из плана группы
    #не удаляет старые Lesson-строки (это история) — отчёт продолжал показывать
    #отменённые дисциплины со старыми оценками, а свежедобавленный предмет без единого
    #занятия не показывался вовсе. Тот же фильтр, что уже чинил это в статистике
    #студента и у Вектора (webdata.current_subject_lessons) — АРХИВ (явно выбранный
    #прошлый термин) не фильтруем: тогда велись другие предметы, прятать их — переписывать
    #прошлое задним числом.
    is_archive = (year, semester) != W.current_term(cfg)
    lessons = W.current_subject_lessons(db, group, lessons, is_archive)
    cutoff = parse_ddmmyyyy(cutoff_date) if cutoff_date else None
    if cutoff:
        # Занятие без даты/с нераспознанной датой — оставляем (лучше лишнее, чем
        # молча потерять данные из-за формата), как и в других местах приложения.
        lessons = [l for l in lessons if (parse_ddmmyyyy(l.date) or cutoff) <= cutoff]
    subjects = sorted({l.subject for l in lessons if l.subject})
    #Группа собирает НЕСКОЛЬКО предметов сразу — у каждого может быть свой преподаватель
    #со своей шкалой (§ролей, 3.3.1), поэтому карта по занятиям, а не одна строка.
    scale_map = W.lesson_scale_map(db, lessons)

    rows = []
    for s in students:
        recs = W.student_records(db, s.surname, s.name, group)
        per_subj = {d["subject"]: d["average"] for d in
                    W.per_subject_averages(lessons, recs, cfg, scale=scale_map)}
        absc = W.absences(lessons, recs)
        rows.append({
            "surname": s.surname, "name": s.name,
            "subject_avg": per_subj,                       # {предмет: средний}
            "average": W.average(lessons, recs, cfg, scale=scale_map),      # общий средний
            "absences": absc,                              # {Н,Б,О,всего}
            "debts": len(W.debts(lessons, recs, scale=scale_map)),
        })
    # средний по группе — по студентам с ненулевым средним (пустые не занижают)
    vals = [r["average"] for r in rows if r["average"]]
    group_avg = round(sum(vals) / len(vals), 2) if vals else 0.0
    # средний ПО КАЖДОМУ ПРЕДМЕТУ (по студентам с ненулевым баллом) — видно, где группа
    # слабее: куратору это нужнее общего числа.
    subject_group_avg = {}
    for subj in subjects:
        sv = [r["subject_avg"].get(subj) for r in rows]
        sv = [x for x in sv if x]
        subject_group_avg[subj] = round(sum(sv) / len(sv), 2) if sv else 0.0
    return {"group": group, "subjects": subjects, "rows": rows,
            "group_avg": group_avg, "subject_group_avg": subject_group_avg,
            "students": len(students)}


#Пороги 4.5/3.5/2.5 — ТЕ ЖЕ, что уже использует остальное приложение для цвета/настроения
#по среднему баллу (vector/mascot.py, ui/dashboards.py, ui/teacher_dashboard.py). Свой
#набор порогов здесь НЕ придумываем — иначе одна и та же оценка красилась бы по-разному
#на разных экранах (§12 отчёта куратора в мессенджере).
def categorize(rows: list) -> dict:
    """Отличники/хорошисты/успевающие/неуспевающие — по общему среднему баллу студента.
    Студентов без единой оценки (average == 0) не считаем ни в одну категорию: «нет
    данных» — это не то же самое, что «неуспевающий»."""
    buckets = {"excellent": 0, "good": 0, "passing": 0, "failing": 0}
    counted = 0
    for r in rows:
        avg = r.get("average") or 0
        if not avg:
            continue
        counted += 1
        if avg >= 4.5:
            buckets["excellent"] += 1
        elif avg >= 3.5:
            buckets["good"] += 1
        elif avg >= 2.5:
            buckets["passing"] += 1
        else:
            buckets["failing"] += 1
    return {"counts": buckets, "total": counted, "no_data": len(rows) - counted}


def _fmt(v) -> str:
    return f"{v:.2f}" if v else "—"


#  Excel

def build_xlsx(groups_data: list, term: dict) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    wb.remove(wb.active)
    thin = Side(style="thin", color="D0D7DE")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    head_fill = PatternFill("solid", fgColor="147C8B")
    head_font = Font(bold=True, color="FFFFFF", size=10)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for gd in groups_data:
        # имя листа: Excel не любит /\?*[] и >31 символа
        safe = gd["group"]
        for ch in "/\\?*[]:":
            safe = safe.replace(ch, "-")
        ws = wb.create_sheet(safe[:31] or "Группа")

        subs = gd["subjects"]
        # шапка: № | Студент | <предметы...> | Общий | Н | Б | О | Пропуски | Долги
        headers = ["№", "Студент"] + subs + ["Общий", "Н", "Б", "О", "Проп.", "Долги"]
        ws.append([f'Группа {gd["group"]} · {term["year"]} · '
                   f'{"осенний" if term["semester"] == 1 else "весенний"} семестр'])
        ws.append([])
        ws.append(headers)
        hr = ws.max_row
        for c in range(1, len(headers) + 1):
            cell = ws.cell(hr, c)
            cell.fill = head_fill
            cell.font = head_font
            cell.alignment = center
            cell.border = border

        for i, r in enumerate(gd["rows"], 1):
            line = [i, f'{r["surname"]} {r["name"]}']
            line += [_fmt(r["subject_avg"].get(s)) for s in subs]
            line += [_fmt(r["average"]), r["absences"]["Н"], r["absences"]["Б"],
                     r["absences"]["О"], r["absences"]["всего"], r["debts"]]
            ws.append(line)
            for c in range(1, len(headers) + 1):
                cell = ws.cell(ws.max_row, c)
                cell.border = border
                if c >= 3:
                    cell.alignment = center

        # строка «Средний по группе»: средний по КАЖДОМУ предмету + общий
        sga = gd.get("subject_group_avg", {})
        ws.append(["", "Средний по группе"] + [_fmt(sga.get(s)) for s in subs]
                  + [_fmt(gd["group_avg"]), "", "", "", "", ""])
        fr = ws.max_row
        for c in range(1, len(headers) + 1):
            cell = ws.cell(fr, c)
            cell.font = Font(bold=True)
            cell.fill = PatternFill("solid", fgColor="E8F0F1")
            if c >= 3:
                cell.alignment = center

        # ширины
        ws.column_dimensions["A"].width = 4
        ws.column_dimensions["B"].width = 26
        for idx in range(len(subs)):
            ws.column_dimensions[get_column_letter(3 + idx)].width = 14
        for idx in range(len(subs) + 3, len(headers)):
            ws.column_dimensions[get_column_letter(idx)].width = 7
        ws.freeze_panes = "C4"

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


#  Word

def build_docx(groups_data: list, term: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT

    doc = Document()
    # альбомная: предметов может быть много, в книжную не влезет
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width

    for gi, gd in enumerate(groups_data):
        if gi:
            doc.add_page_break()
        term_word = "осенний" if term["semester"] == 1 else "весенний"
        h = doc.add_heading(f'Успеваемость группы {gd["group"]}', level=1)
        sub = doc.add_paragraph(f'{term["year"]} · {term_word} семестр · '
                                f'студентов: {gd["students"]} · '
                                f'средний по группе: {_fmt(gd["group_avg"])}')
        sub.runs[0].italic = True

        subs = gd["subjects"]
        headers = ["№", "Студент"] + subs + ["Общий", "Н", "Б", "О", "Долги"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Light Grid Accent 1"
        for c, txt in enumerate(headers):
            run = table.rows[0].cells[c].paragraphs[0].add_run(txt)
            run.bold = True
            run.font.size = Pt(8)

        for i, r in enumerate(gd["rows"], 1):
            cells = table.add_row().cells
            vals = [str(i), f'{r["surname"]} {r["name"]}']
            vals += [_fmt(r["subject_avg"].get(s)) for s in subs]
            vals += [_fmt(r["average"]), str(r["absences"]["Н"]),
                     str(r["absences"]["Б"]), str(r["absences"]["О"]), str(r["debts"])]
            for c, v in enumerate(vals):
                p = cells[c].paragraphs[0]
                run = p.add_run(v)
                run.font.size = Pt(8)
                if c >= 2:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # итоговая строка «Средний по группе»: по каждому предмету + общий
        sga = gd.get("subject_group_avg", {})
        foot = table.add_row().cells
        fvals = ["", "Средний по группе"] + [_fmt(sga.get(s)) for s in subs] \
            + [_fmt(gd["group_avg"]), "", "", "", ""]
        for c, v in enumerate(fvals):
            p = foot[c].paragraphs[0]
            run = p.add_run(v)
            run.font.size = Pt(8)
            run.bold = True
            if c >= 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
